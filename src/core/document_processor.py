"""Обработка документов с паттерном «Фабрика» для расширяемости."""

from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path

from utils.exceptions import DocumentParsingError, UnsupportedFormatError
from utils.logging_config import get_logger

logger = get_logger("core.document_processor")


class DocumentProcessor(ABC):
    """Абстрактная база для извлечения текста из документов."""

    @abstractmethod
    def extract_text(self, path: Path) -> str:
        """Извлечь обычный текст из файла документа.

        Аргументы:
            path: Путь к документу.

        Возвращает:
            Извлечённый текст одной строкой.

        Исключения:
            DocumentParsingError: при невозможности разбора файла.
        """

    @abstractmethod
    def supported_extensions(self) -> tuple[str, ...]:
        """Расширения файлов, обрабатываемые этим процессором (напр. (.docx,))."""


class DocxProcessor(DocumentProcessor):
    """Извлечение текста из файлов Microsoft Word .docx с помощью python-docx."""

    def extract_text(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError as exc:
            raise DocumentParsingError(
                str(path), "python-docx is not installed"
            ) from exc

        if not path.exists():
            raise DocumentParsingError(str(path), "File does not exist")

        try:
            doc = Document(str(path))
            paragraphs: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Также извлечь текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)

            full_text = "\n".join(paragraphs)
            logger.info(
                "Extracted %d characters from %s (%d paragraphs)",
                len(full_text), path.name, len(paragraphs),
            )
            return full_text
        except Exception as exc:
            raise DocumentParsingError(str(path), str(exc)) from exc

    def supported_extensions(self) -> tuple[str, ...]:
        return (".docx",)


class DocumentProcessorFactory:
    """Фабрика, возвращающая подходящий DocumentProcessor для файла.

    Расширяемо: новые процессоры регистрируются через register.
    """

    _processors: dict[str, DocumentProcessor] = {}

    @classmethod
    def register(cls, processor: DocumentProcessor) -> None:
        """Зарегистрировать процессор для поддерживаемых им расширений."""
        for ext in processor.supported_extensions():
            cls._processors[ext.lower()] = processor
            logger.debug("Registered processor for '%s'", ext)

    @classmethod
    def create(cls, path: Path) -> DocumentProcessor:
        """Получить процессор по расширению файла.

        Аргументы:
            path: Путь к файлу документа.

        Возвращает:
            Подходящий DocumentProcessor.

        Исключения:
            UnsupportedFormatError: если процессор не зарегистрирован.
        """
        ext = path.suffix.lower()
        processor = cls._processors.get(ext)
        if processor is None:
            raise UnsupportedFormatError(str(path), ext)
        return processor

    @classmethod
    def register_defaults(cls) -> None:
        """Зарегистрировать встроенные процессоры."""
        cls.register(DocxProcessor())
